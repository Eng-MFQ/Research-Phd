"""
Convert Scientific_Report_Expanded.md to a .docx that complies with the
"Methodological Recommendations for Preparing a Scientific Report":

  - Font: Times New Roman, 14 pt
  - Line spacing: 1.5
  - Margins: left 3 cm, right 1.5 cm, top 2 cm, bottom 2.5 cm
"""

import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


HERE = Path(__file__).parent
SRC = HERE / "Scientific_Report_Expanded.md"
DST = HERE / "Scientific_Report_Expanded.docx"
TITLE_TEMPLATE = (
    HERE.parent
    / "Phd Important files"
    / "Attachments"
    / "5 TEMPLATE (REPORT, en).docx"
)
# Number of (paragraph + table) body elements from the template that
# constitute the official title page, ending at "Иннополис / Innopolis, 2026".
# In the template, body indices 0..29 inclusive cover the title page, but one
# of those (index 16) is a bookmarkEnd that we skip, leaving 29 p/tbl elements.
TITLE_PAGE_ELEMENTS = 29


# ----- Document setup -----------------------------------------------------

def setup_document() -> Document:
    doc = Document()

    # Page margins.
    for section in doc.sections:
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.5)

    # Default body style: TNR 14pt, 1.5 line spacing.
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    # Ensure East-Asian / complex-script fallback also TNR (Word quirk).
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), "Times New Roman")

    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)

    # Heading styles: same TNR family, 1.5 spacing, bold, scaled sizes.
    heading_sizes = {1: 18, 2: 16, 3: 14, 4: 14, 5: 14}
    for level, size in heading_sizes.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True
        # Force TNR for cs/eastAsia too.
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), "Times New Roman")

    return doc


# ----- Inline parser ------------------------------------------------------

INLINE_RE = re.compile(
    r"(\*\*(?P<bold>[^*]+)\*\*)"        # **bold**
    r"|(\*(?P<italic>[^*]+)\*)"          # *italic*
    r"|(`(?P<code>[^`]+)`)"              # `code`
    r"|(\$\$(?P<dmath>.+?)\$\$)"        # $$display math$$
    r"|(\$(?P<imath>[^$]+)\$)"           # $inline math$
)


def add_inline_runs(paragraph, text: str) -> None:
    """Apply bold/italic/code/math formatting inside a paragraph."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group("bold") is not None:
            run = paragraph.add_run(m.group("bold"))
            run.bold = True
        elif m.group("italic") is not None:
            run = paragraph.add_run(m.group("italic"))
            run.italic = True
        elif m.group("code") is not None:
            run = paragraph.add_run(m.group("code"))
            run.font.name = "Courier New"
            run.font.size = Pt(12)
        elif m.group("dmath") is not None:
            run = paragraph.add_run(m.group("dmath"))
            run.italic = True
        elif m.group("imath") is not None:
            run = paragraph.add_run(m.group("imath"))
            run.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


# ----- Block parser -------------------------------------------------------

HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
ULIST_RE = re.compile(r"^(\s*)[*\-]\s+(.*)$")
OLIST_RE = re.compile(r"^(\s*)(\d+)\.\s+(.*)$")
TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$")


def is_table_row(line: str) -> bool:
    return line.lstrip().startswith("|") and line.rstrip().endswith("|")


def parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [c.strip() for c in line.split("|")]


def render(doc: Document, lines: list[str]) -> None:
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip blank lines.
        if not stripped:
            i += 1
            continue

        # Horizontal rule.
        if stripped == "---":
            i += 1
            continue

        # Code fence.
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # consume closing fence
            for cl in code_lines:
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(cl)
                run.font.name = "Courier New"
                run.font.size = Pt(11)
            doc.add_paragraph()  # spacer
            continue

        # Headings.
        m = HEADING_RE.match(stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            heading = doc.add_paragraph(style=f"Heading {min(level, 5)}")
            add_inline_runs(heading, text)
            i += 1
            continue

        # Table (header row followed by separator row).
        if is_table_row(line) and i + 1 < len(lines) and TABLE_SEP_RE.match(lines[i + 1]):
            header = parse_table_row(line)
            i += 2  # skip header + separator
            body = []
            while i < len(lines) and is_table_row(lines[i]):
                body.append(parse_table_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1 + len(body), cols=len(header))
            table.style = "Light Grid Accent 1"
            for c, h in enumerate(header):
                cell = table.rows[0].cells[c]
                cell.text = ""
                p = cell.paragraphs[0]
                add_inline_runs(p, h)
                for run in p.runs:
                    run.bold = True
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
            for r, row in enumerate(body, start=1):
                for c, val in enumerate(row):
                    cell = table.rows[r].cells[c]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    add_inline_runs(p, val)
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)
            doc.add_paragraph()
            continue

        # Unordered list.
        m = ULIST_RE.match(line)
        if m:
            while i < len(lines):
                m2 = ULIST_RE.match(lines[i])
                if not m2:
                    break
                p = doc.add_paragraph(style="List Bullet")
                add_inline_runs(p, m2.group(2))
                i += 1
            continue

        # Ordered list.
        m = OLIST_RE.match(line)
        if m:
            while i < len(lines):
                m2 = OLIST_RE.match(lines[i])
                if not m2:
                    break
                p = doc.add_paragraph(style="List Number")
                add_inline_runs(p, m2.group(3))
                i += 1
            continue

        # Display math block on its own line.
        if stripped.startswith("$$") and stripped.endswith("$$") and len(stripped) > 4:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped.strip("$"))
            run.italic = True
            i += 1
            continue

        # Plain paragraph: collect contiguous non-blank, non-special lines.
        buf = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i]
            nxt_stripped = nxt.strip()
            if (
                not nxt_stripped
                or HEADING_RE.match(nxt_stripped)
                or ULIST_RE.match(nxt)
                or OLIST_RE.match(nxt)
                or is_table_row(nxt)
                or nxt_stripped.startswith("```")
                or nxt_stripped == "---"
            ):
                break
            buf.append(nxt_stripped)
            i += 1
        p = doc.add_paragraph()
        add_inline_runs(p, " ".join(buf))


# ----- Title page prepend -------------------------------------------------

def prepend_title_page(doc: Document, template_path: Path, n_elements: int) -> None:
    """Insert the first `n_elements` body elements (paragraphs + tables) from
    `template_path` at the start of `doc`, followed by a hard page break so
    the rendered report begins on the next page."""
    template = Document(str(template_path))
    template_body = template.element.body
    target_body = doc.element.body

    # The current first child of the target's body — we insert before it.
    target_first = target_body[0]

    # Collect the title-page elements (skip bookmarkStart/End markers that
    # python-docx attaches around named bookmarks; they reference IDs that
    # would dangle in the new doc).
    collected = []
    seen = 0
    for child in list(template_body):
        tag = child.tag.split("}")[-1]
        if tag in ("p", "tbl"):
            if seen >= n_elements:
                break
            collected.append(deepcopy(child))
            seen += 1

    # Strip bookmark markers from inside the copied content to avoid ID clashes.
    for elem in collected:
        for bm in elem.findall(".//" + qn("w:bookmarkStart")):
            bm.getparent().remove(bm)
        for bm in elem.findall(".//" + qn("w:bookmarkEnd")):
            bm.getparent().remove(bm)

    # Hard page break before the rendered content.
    page_break_p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    page_break_p.append(r)
    collected.append(page_break_p)

    for elem in collected:
        target_first.addprevious(elem)


# ----- Main ---------------------------------------------------------------

def main() -> None:
    md = SRC.read_text(encoding="utf-8")
    lines = md.splitlines()

    doc = setup_document()
    render(doc, lines)

    if TITLE_TEMPLATE.exists():
        prepend_title_page(doc, TITLE_TEMPLATE, TITLE_PAGE_ELEMENTS)
    else:
        print(f"Warning: title-page template not found at {TITLE_TEMPLATE}")

    doc.save(DST)
    print(f"Wrote {DST}")


if __name__ == "__main__":
    main()
