"""
Post-process Scientific_Report_Expanded.docx to satisfy the formatting
requirements of the Scientific Report:

  1. Times New Roman, 14 pt, 1.5 line spacing, justified body text.
  2. Page margins: left 3 cm, right 1.5 cm, top 2 cm, bottom 2.5 cm.
  3. Page numbers in the footer (no number on the title page).
  4. Each ordered (numbered) list restarts at 1 instead of continuing the
     count across the whole document (the bug that produced 6-9, 38-40 ...).

The script edits the document in place; a *.backup.docx copy is kept separately.
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = Path(__file__).parent
DOC = HERE / "Scientific_Report_Expanded.docx"

# Style of the numbered-list paragraphs and the abstract numbering they use.
LIST_NUMBER_STYLE = "List Number"
LIST_NUMBER_ABSTRACT_ID = "7"   # abstractNumId backing the "List Number" style


# --------------------------------------------------------------------------
# 1. Base formatting (font / size / spacing / margins)
# --------------------------------------------------------------------------

def enforce_base_formatting(doc: Document) -> None:
    for section in doc.sections:
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(14)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), "Times New Roman")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


# --------------------------------------------------------------------------
# 2. Justify body text (leave centered title-page / heading content alone)
# --------------------------------------------------------------------------

JUSTIFY_STYLES = {"Normal", "List Bullet", "List Number",
                  "List Bullet 2", "List Number 2"}


def justify_body(doc: Document) -> None:
    for p in doc.paragraphs:
        if p.style.name not in JUSTIFY_STYLES:
            continue
        # Preserve intentionally centered/right content (e.g. title page).
        if p.alignment in (WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT):
            continue
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


# --------------------------------------------------------------------------
# 3. Page numbers in the footer (title page excluded)
# --------------------------------------------------------------------------

def _page_number_paragraph(footer):
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r = run._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_end)


def add_page_numbers(doc: Document) -> None:
    section = doc.sections[0]
    # Different first page so the title page carries no number.
    section.different_first_page_header_footer = True
    sectPr = section._sectPr
    titlePg = sectPr.find(qn("w:titlePg"))
    if titlePg is None:
        titlePg = OxmlElement("w:titlePg")
        sectPr.append(titlePg)

    footer = section.footer
    footer.is_linked_to_previous = False
    _page_number_paragraph(footer)

    # Ensure the first-page footer exists but stays empty (no number).
    first_footer = section.first_page_footer
    first_footer.is_linked_to_previous = False


# --------------------------------------------------------------------------
# 4. Restart numbering for every ordered list
# --------------------------------------------------------------------------

def _numbering_root(doc: Document):
    return doc.part.numbering_part.numbering_definitions._numbering


def _max_num_id(numbering) -> int:
    ids = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    return max(ids) if ids else 0


def _make_num(numbering, num_id: int, abstract_id: str):
    """Create a <w:num> that reuses abstract_id but restarts at 1."""
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs = OxmlElement("w:abstractNumId")
    abs.set(qn("w:val"), abstract_id)
    num.append(abs)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    # <w:num> elements must follow all <w:abstractNum> elements.
    numbering.append(num)
    return num


def _set_paragraph_num(paragraph, num_id: int) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        numPr = OxmlElement("w:numPr")
        pPr.append(numPr)
    ilvl = numPr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        numPr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    numId = numPr.find(qn("w:numId"))
    if numId is None:
        numId = OxmlElement("w:numId")
        numPr.append(numId)
    numId.set(qn("w:val"), str(num_id))


def restart_ordered_lists(doc: Document) -> int:
    numbering = _numbering_root(doc)
    next_id = _max_num_id(numbering) + 1

    lists_found = 0
    in_list = False
    current_id = None
    for p in doc.paragraphs:
        if p.style.name == LIST_NUMBER_STYLE:
            if not in_list:
                # Start of a new ordered list -> fresh restarting numId.
                current_id = next_id
                next_id += 1
                _make_num(numbering, current_id, LIST_NUMBER_ABSTRACT_ID)
                lists_found += 1
                in_list = True
            _set_paragraph_num(p, current_id)
        else:
            in_list = False
            current_id = None
    return lists_found


# --------------------------------------------------------------------------

def main() -> None:
    doc = Document(str(DOC))
    enforce_base_formatting(doc)
    justify_body(doc)
    add_page_numbers(doc)
    n = restart_ordered_lists(doc)
    doc.save(str(DOC))
    print(f"Saved {DOC}")
    print(f"Restarted numbering for {n} ordered lists.")


if __name__ == "__main__":
    main()
